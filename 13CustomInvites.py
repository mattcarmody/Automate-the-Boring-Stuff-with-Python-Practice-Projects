import docx

guest_list = open(FILE_HERE).read().splitlines()
doc = docx.Document('/home/matt/AutomateBookLocal/invitations.docx')

for guest in guest_list:
	doc.add_paragraph('It would be a pleasure to have the company of', style='inviteCurly')
	doc.add_paragraph(guest, style='inviteNorm')
	doc.add_paragraph('at Memory Lane on the evening of', style='inviteCurly')
	doc.add_paragraph('April 1st', style='inviteNorm')
	doc.add_paragraph("at 7 o'clock", style="inviteCurly")
	doc.add_page_break()
	
doc.save(FILE_HERE)
