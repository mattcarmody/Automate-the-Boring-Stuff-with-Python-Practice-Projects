#! /usr/bin/python3
# chap13PracProjCustomInvites.py

import docx

guestFile = open('/home/matt/AutomateBook/automate_online-materials/guests.txt')
guestList = guestFile.read().splitlines()
doc = docx.Document('/home/matt/AutomateBook/invitations.docx')

for guest in guestList:
	doc.add_paragraph('It would be a pleasure to have the company of', style='inviteCurly')
	doc.add_paragraph(guest, style='inviteNorm')
	doc.add_paragraph('at Memory Lane on the evening of', style='inviteCurly')
	doc.add_paragraph('April 1st', style='inviteNorm')
	doc.add_paragraph("at 7 o'clock", style="inviteCurly")
	doc.add_page_break()
	
doc.save('/home/matt/AutomateBook/invitations2.docx')
